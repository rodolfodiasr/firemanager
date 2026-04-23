"""Document generation service — PDF audit report and DOCX manual guide."""
import io
from datetime import datetime, timezone
from pathlib import Path
from uuid import UUID

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document, DocumentType
from app.models.operation import Operation

_DOCS_DIR = Path("/tmp/firemanager_docs")
_DOCS_DIR.mkdir(parents=True, exist_ok=True)


async def generate_documents_for_operation(db: AsyncSession, operation_id: UUID) -> list[Document]:
    result = await db.execute(select(Operation).where(Operation.id == operation_id))
    operation = result.scalar_one_or_none()
    if not operation:
        return []

    docs = []
    docs.append(await _generate_audit_pdf(db, operation))
    docs.append(await _generate_manual_docx(db, operation))
    return [d for d in docs if d is not None]


async def _generate_audit_pdf(db: AsyncSession, operation: Operation) -> Document | None:
    try:
        from weasyprint import HTML

        html_content = _build_audit_html(operation)
        filename = f"audit_report_{operation.id}.pdf"
        path = _DOCS_DIR / filename
        HTML(string=html_content).write_pdf(str(path))

        doc = Document(
            operation_id=operation.id,
            doc_type=DocumentType.audit_report_pdf,
            filename=filename,
            storage_path=str(path),
        )
        db.add(doc)
        await db.flush()
        return doc
    except Exception:
        return None


async def _generate_manual_docx(db: AsyncSession, operation: Operation) -> Document | None:
    try:
        from docx import Document as DocxDocument

        docx = DocxDocument()
        docx.add_heading("FireManager — Guia de Configuração Manual", 0)
        docx.add_paragraph(f"Data: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
        docx.add_heading("Operação", level=1)
        docx.add_paragraph(f"ID: {operation.id}")
        docx.add_paragraph(f"Solicitação: {operation.natural_language_input}")
        docx.add_paragraph(f"Status: {operation.status.value}")

        if operation.action_plan:
            docx.add_heading("Plano de Ação", level=1)
            import json
            docx.add_paragraph(json.dumps(operation.action_plan, indent=2, ensure_ascii=False))

        filename = f"manual_guide_{operation.id}.docx"
        path = _DOCS_DIR / filename
        docx.save(str(path))

        doc = Document(
            operation_id=operation.id,
            doc_type=DocumentType.manual_guide_docx,
            filename=filename,
            storage_path=str(path),
        )
        db.add(doc)
        await db.flush()
        return doc
    except Exception:
        return None


def _build_audit_html(operation: Operation) -> str:
    import json
    plan_json = json.dumps(operation.action_plan or {}, indent=2, ensure_ascii=False)
    return f"""
<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="UTF-8">
  <title>FireManager — Relatório de Auditoria</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 40px; }}
    h1 {{ color: #e63946; }}
    pre {{ background: #f4f4f4; padding: 16px; border-radius: 4px; }}
    .meta {{ color: #666; font-size: 0.9em; }}
  </style>
</head>
<body>
  <h1>🔥 FireManager — Relatório de Auditoria</h1>
  <p class="meta">Gerado em {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}</p>
  <h2>Operação {operation.id}</h2>
  <p><strong>Solicitação original:</strong> {operation.natural_language_input}</p>
  <p><strong>Intenção identificada:</strong> {operation.intent or 'N/A'}</p>
  <p><strong>Status:</strong> {operation.status.value}</p>
  <h2>Plano de Ação</h2>
  <pre>{plan_json}</pre>
</body>
</html>
"""
